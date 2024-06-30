from math import ceil

from bson import ObjectId
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt

from config.mongo_config import answers, questions, themes, users


def set_cell_color(cell, color):
    cell_xml_element = cell._tc
    table_cell_properties = cell_xml_element.get_or_add_tcPr()
    shade_obj = OxmlElement('w:shd')
    shade_obj.set(qn('w:fill'), color)
    table_cell_properties.append(shade_obj)


def create_tests_docx_file(plan):
    department = plan.get('department')
    quarter = plan.get('quarter')
    year = plan.get('year')
    DOC_HEADER = (
        'Тестовое задание\n'
        'по входному и выходному контролю знаний технической учёбы\n'
        f'{department} на {quarter} квартал {year}г.'
    )
    TABLE_HEADERS = (
        '№ п/п',
        'Тема технической учёбы',
        '№ вопроса',
        'Содержание вопроса',
        'Варианты ответов',
        '',
    )
    document = Document()
    section = document.sections[0]
    # левое поле в миллиметрах
    section.left_margin = Mm(25)
    # правое поле в миллиметрах
    section.right_margin = Mm(15)
    # верхнее поле в миллиметрах
    section.top_margin = Mm(15)
    # нижнее поле в миллиметрах
    section.bottom_margin = Mm(15)

    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    header = document.add_paragraph(DOC_HEADER)
    header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    header.style = document.styles['Normal']

    table = document.add_table(rows=2, cols=6, style='Table Grid')
    table_hdr = table.rows[0].cells  # заголовки таблицы
    table_num = table.rows[1].cells  # нумерация таблицы
    table.columns[0].width = Mm(10)
    table.columns[1].width = Mm(30)
    table.columns[2].width = Mm(20)
    table.columns[3].width = Mm(53)
    table.columns[4].width = Mm(7)
    table.columns[5].width = Mm(58)

    for id, cell in enumerate(table_hdr):
        cell.text = TABLE_HEADERS[id]
        table_num[id].text = str(id + 1)

    table_hdr[4].merge(table_hdr[5])

    questions_list = plan.get('questions')

    current_rows = 2  # текущее количество строк в таблице
    for num, q_id in enumerate(questions_list):
        q = questions.find_one({'_id': ObjectId(q_id)})
        q_answers = list(answers.find({'q_id': ObjectId(q_id), 'is_active': True}))
        num_answers = len(q_answers)
        main_row = table.add_row().cells
        main_row[0].text = str(num + 1)
        main_row[1].text = themes.find_one({'code': q.get('theme')}).get('name')
        main_row[2].text = str(num + 1)
        main_row[3].text = q.get('text')
        main_row[4].text = '1'
        main_row[5].text = q_answers[0]['text']
        for i in range(1, num_answers):
            row_ans = table.add_row().cells
            row_ans[-2].text = str(i + 1)
            row_ans[-1].text = q_answers[i]['text']
        for row_id in range(current_rows, current_rows + num_answers - 1):
            for cell_id in range(0, 4):
                table.rows[row_id].cells[cell_id].merge(table.rows[row_id+1].cells[cell_id])
        current_rows += num_answers

    path = f'static/export/Тест {department} ({quarter} кв. {year}г).docx'
    document.save(path)


def create_results_docx_file(year, quarter, test_type, results_set):
    DOC_HEADER = (
        f'Результаты тестирования {test_type} контроля знаний по технической учебе\n'
        f'за {quarter} квартал {year} года'
    )
    TABLE_HEADERS = ('№ вопроса', 'Вариант ответа')
    document = Document()
    section = document.sections[0]
    section.left_margin = Mm(25)
    section.right_margin = Mm(15)
    section.top_margin = Mm(15)
    section.bottom_margin = Mm(15)

    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    for res in results_set:
        results = res.get('quiz_results')
        username = users.find_one({'user_id': res.get('user_id')}).get('username')
        grade = res.get('grade')

        len_test = len(results['answers'])
        num_rows = ceil(len_test / 10) * 2
        header = document.add_paragraph(DOC_HEADER)
        header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        header.style = document.styles['Normal']

        table = document.add_table(rows=num_rows, cols=11, style='Table Grid')
        table.columns[0].width = Mm(25)
        for c in range(1, 11):
            table.columns[c].width = Mm(15)

        table_hdr = table.columns[0].cells
        for id, cell in enumerate(table_hdr):
            cell_id = id % 2  # определение нечётных полей
            cell.text = TABLE_HEADERS[cell_id]
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        for id, ans_id in enumerate(results['answers']):
            if ans_id is not None:
                ans = answers.find_one({'_id': ObjectId(ans_id)})
                q_answers = list(answers.find({'q_id': ans['q_id'], 'is_active': True}))
                for a_id, ans_data in enumerate(q_answers):
                    if str(ans_data['_id']) == ans_id:
                        ans_pos = (a_id + 1)
                        break
                correct = ans['is_correct']
            else:
                ans_pos = '-'
                correct = False

            row_id = (id // 10) * 2
            cell_id = int(str(id)[-1])
            cell_hdr = table.rows[row_id].cells[cell_id+1]
            cell_value = table.rows[row_id+1].cells[cell_id+1]
            cell_hdr.text = str(id + 1)
            cell_hdr.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            cell_value.text = str(ans_pos)
            cell_value.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            if correct:
                set_cell_color(cell_value, 'addfad')

        document.add_paragraph('')
        document.add_paragraph(f'Тестируемый   ______________   {username}')
        document.add_paragraph(f'Результат тестирования: {grade}')
        document.add_paragraph('')
        document.add_paragraph('')

        path = f'static/export/Результаты {test_type} контроля знаний ({quarter} кв. {year}г).docx'
        document.save(path)
