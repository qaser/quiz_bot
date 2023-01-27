from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

DOC_HEADER = 'Тестовое задание\nпо входному и выходному контролю знаний технической учёбы\nслужбы ГКС на 1 квартал 2023г.'

TABLE_HEADERS = (
    '№ п/п',
    'Тема технической учёбы',
    '№ вопроса',
    'Содержание вопроса',
    'Варианты ответов',
)

document = Document()
style = document.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

header = document.add_paragraph(DOC_HEADER)
header.alignment = WD_ALIGN_PARAGRAPH.CENTER
header.style = document.styles['Normal']

table = document.add_table(rows=3, cols=6)

# a = table.cell(0, 0)
# b = table.cell(1, 1)
# c = table.cell(2, 1)
# d = table.cell(2, 2)

# A = a.merge(b)
# B = A.merge(c)
# C = B.merge(d)

document.save('demo.docx')
